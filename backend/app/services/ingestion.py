from __future__ import annotations
import logging
from pathlib import Path
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from app.core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
CHUNK_SIZE = 800
CHUNK_OVERLAP = 120

_embeddings = None
_store = None


def get_store():
    global _embeddings, _store
    if _store is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
        persist_dir = Path(settings.chroma_persist_dir) / "default"
        persist_dir.mkdir(parents=True, exist_ok=True)
        _store = Chroma(
            collection_name="smartchat_default",
            embedding_function=_embeddings,
            persist_directory=str(persist_dir),
        )
    return _store


def _load_document(file_path: Path) -> list[Document]:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        loader = PyPDFLoader(str(file_path))
        docs = loader.load()
    elif ext == ".docx":
        from docx import Document as DocxDocument
        docx = DocxDocument(str(file_path))
        content = "\n\n".join(
            para.text for para in docx.paragraphs if para.text.strip()
        )
        docs = [Document(page_content=content,
                         metadata={"source_file": file_path.name})]
    elif ext in (".txt", ".md"):
        content = file_path.read_text(encoding="utf-8")
        docs = [Document(page_content=content,
                         metadata={"source_file": file_path.name})]
    else:
        raise ValueError(f"Tipo no soportado: {ext}")
    for doc in docs:
        doc.metadata["source_file"] = file_path.name
    return docs


def _split_documents(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(docs)


async def ingest_files(file_paths: list[Path]) -> tuple[int, list[str]]:
    all_chunks: list[Document] = []
    processed: list[str] = []

    for path in file_paths:
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            docs = _load_document(path)
            chunks = _split_documents(docs)
            all_chunks.extend(chunks)
            processed.append(path.name)
        except Exception as e:
            logger.error(f"Error loading {path.name}: {e}")

    if not all_chunks:
        raise ValueError("No se pudo procesar ningún documento válido.")

    store = get_store()
    store.add_documents(all_chunks)
    return len(all_chunks), processed
